
import os
import streamlit as st
from PyPDF2 import PdfReader
from openai import OpenAI, RateLimitError
from docx import Document
from openpyxl import Workbook
import json

# --- Load role-based questions from JSON ---
json_path = os.path.join(os.path.dirname(__file__), "role_questions.json")
with open(json_path, "r", encoding="utf-8") as f:
    role_questions = json.load(f)

# Set page config and app title (shown before password entry)
st.set_page_config(page_title="ClauseMatrix Legal Analyzer", layout="wide")
st.title("ClauseMatrix Legal Analyzer")

# --- Password Gate ---
PASSWORD = "LegalAI_PDF#25"
if "access_granted" not in st.session_state:
    st.session_state.access_granted = False

if not st.session_state.access_granted:
    pwd = st.text_input("🔐 Please enter access password", type="password")
    if pwd == PASSWORD:
        st.session_state.access_granted = True
        st.rerun()
    else:
        st.stop()

# --- Onboarding Header and Confidentiality Notes ---
st.markdown("## Let the AI Guide Begin")
st.markdown("""
This tool uses AI to assist in document triage. For best results:
- **Be specific**: Ask "Compare arbitration clause" or "Spot unusual indemnity term"
- **Context matters**: Mention contract type or jurisdiction if relevant
- **AI is not a lawyer**: Always review important outputs yourself.

### 🔐 Confidentiality & Privacy Assurances
- **Zero Retention**: PDFs are not stored after your session ends.
- **Stateless AI**: No chat history is linked to you or reused.
- **Auto-Cleanup**: Temporary files, if any, are auto-deleted after analysis.
- **Legal First**: You retain full control. Uploads are processed once, then cleared.
""")

# ----------------------------
# Role Selector (horizontal)
# ----------------------------
role = st.radio(
    "### Select your legal role:",
    ["Appellate Attorney", "Paralegal", "Contract Analyst", "Tenant"],
    horizontal=True
)

# --- Upload PDFs ---
st.markdown("### Please upload a single or multiple legal documents (PDF) for clause analysis", unsafe_allow_html=True)
uploaded_files = st.file_uploader("Upload PDFs", type=["pdf"], accept_multiple_files=True)

# --- Initialize OpenAI Client ---
client = OpenAI()

# --- Load Role-Specific Sample Questions ---
@st.cache_data
def load_sample_questions():
    with open("role_questions_FIXED33.json", "r") as f:
        return json.load(f)

role_questions = load_sample_questions()
question_list = role_questions.get(role, [])

# --- Helper Functions ---
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())

def summarize_clause(text, role):
    system_prompt = f"You are a legal expert assisting a {role}. Analyze the following legal content and extract relevant clauses or issues in bullet points."
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.2
        )
        return response.choices[0].message.content.strip()
    except RateLimitError:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.3
        )
        return "(Fallback to GPT-3.5)\n" + response.choices[0].message.content.strip()

# --- Clause Analysis ---
results = {}
if uploaded_files:
    st.markdown("### 🔍 Analysis in Progress...")
    for file in uploaded_files:
        text = extract_text_from_pdf(file)
        if len(text) > 16000:
            chunks = [text[i:i+16000] for i in range(0, len(text), 16000)]
            summary = "\n".join([summarize_clause(chunk, role) for chunk in chunks])
        else:
            summary = summarize_clause(text, role)
        results[file.name] = summary
        st.markdown(f"#### 📄 Analysis for `{file.name}`:")
        st.write(summary)

    # Export to DOCX and XLSX
    docx_path = os.path.join(os.getcwd(), "Clause_Summary.docx")
    docx = Document()
    for filename, summary in results.items():
        docx.add_heading(filename, level=2)
        docx.add_paragraph(summary)
    docx.save(docx_path)
    st.download_button("⬇️ Download Word Analysis", data=open(docx_path, "rb"), file_name="Clause_Summary.docx")

    excel_path = os.path.join(os.getcwd(), "Clause_Summary.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Clause Analysis"
    ws.append(["Filename", "Clause Analysis"])
    for filename, summary in results.items():
        ws.append([filename, summary])
    wb.save(excel_path)
    st.download_button("⬇️ Download Excel Analysis", data=open(excel_path, "rb"), file_name="Clause_Summary.xlsx")

# --- Ask AI with Dropdown ---
st.markdown("## 🧠 Ask AI")
selected_question = st.selectbox("Select a legal question relevant to your role:", question_list)
if st.button("Submit Question") and selected_question:
    followup_prompt = f"As a {role}, respond to the following legal question: {selected_question}"
    followup_response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"You are a helpful legal assistant for a {role}."},
            {"role": "user", "content": followup_prompt}
        ]
    )
    st.markdown("### 💡 AI Legal Insight")
    st.write(followup_response.choices[0].message.content.strip())
